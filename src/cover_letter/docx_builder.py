"""Render a CoverLetter dataclass to .docx.

Single column, Arial, 0.8" margins. Layout matches the master EM template:
  Header (name + contact)
  Subject line
  Greeting
  Opening hook
  Optional company hook
  Background paragraph
  Lead-in line + 3 signal bullets
  Optional company-fit line
  Sign-off paragraph
  Closing
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from ..resume import profile
from .template import CoverLetter


def _set_arial(run, size: int = 11, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def build_docx(letter: CoverLetter, output_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.25

    # ----- Header: Name (centered, ALL CAPS, letter-spaced) -----
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(profile.NAME.upper())
    name_run.bold = True
    name_run.font.name = "Arial"
    name_run.font.size = Pt(16)
    name_run.font.spacing = Pt(2)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_arial(contact_p.add_run(profile.CONTACT_LINE), size=10)

    # ----- Greeting -----
    # The greeting is the first non-header content (no "Re: ..." subject line —
    # the master template doesn't include one, and the opening hook already
    # announces the role).
    name = letter.hiring_manager_name.strip() if letter.hiring_manager_name else ""
    greeting = f"Dear {name}," if name else "Dear Hiring Manager,"
    greet_p = doc.add_paragraph()
    greet_p.paragraph_format.space_before = Pt(14)
    _set_arial(greet_p.add_run(greeting), size=11)
    greet_p.paragraph_format.space_after = Pt(8)

    # ----- Opening hook -----
    opening_p = doc.add_paragraph()
    _set_arial(opening_p.add_run(letter.opening_hook.strip()), size=11)
    opening_p.paragraph_format.space_after = Pt(8)

    # ----- Company hook (optional) -----
    if letter.has_company_hook:
        hook_p = doc.add_paragraph()
        _set_arial(hook_p.add_run(letter.company_hook.strip()), size=11)
        hook_p.paragraph_format.space_after = Pt(8)

    # ----- Background paragraph (locked, frame-dependent) -----
    background = (
        profile.COVER_LETTER_BACKGROUND_HYBRID
        if (letter.frame or "standard").lower() == "hybrid"
        else profile.COVER_LETTER_BACKGROUND_STANDARD
    )
    bg_p = doc.add_paragraph()
    _set_arial(bg_p.add_run(background), size=11)
    bg_p.paragraph_format.space_after = Pt(8)

    # ----- Bullets lead-in + 3 signal bullets -----
    lead_p = doc.add_paragraph()
    _set_arial(lead_p.add_run(profile.COVER_LETTER_BULLETS_LEAD), size=11)
    lead_p.paragraph_format.space_after = Pt(4)

    for pick in letter.bullets:
        text = (pick.bullet or "").strip()
        if not text:
            continue
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(4)
        _set_arial(bp.add_run(text), size=11)

    # ----- Optional company-fit line -----
    if letter.has_company_fit_line:
        fit_p = doc.add_paragraph()
        _set_arial(fit_p.add_run(letter.company_fit_line.strip()), size=11)
        fit_p.paragraph_format.space_before = Pt(8)
        fit_p.paragraph_format.space_after = Pt(8)

    # ----- Sign-off paragraph (locked) -----
    signoff_p = doc.add_paragraph()
    _set_arial(signoff_p.add_run(profile.COVER_LETTER_SIGNOFF), size=11)
    signoff_p.paragraph_format.space_before = Pt(8)
    signoff_p.paragraph_format.space_after = Pt(8)

    # ----- Closing (locked) -----
    for line in profile.COVER_LETTER_CLOSING.split("\n"):
        p = doc.add_paragraph()
        _set_arial(p.add_run(line), size=11)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
