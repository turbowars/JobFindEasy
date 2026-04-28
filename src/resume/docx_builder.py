"""Render a Resume dataclass to .docx.

Single column, Arial, 0.6" margins. Section order matches the spec:
Header (name + contact) -> Summary -> Highlights -> Experience ->
Education & Certifications -> Skills.

Bullets honor inline `**bold**` markers so the LLM can emphasize metric
phrases without us round-tripping through markdown.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .template import Resume

CONTENT_WIDTH_INCHES = 7.3  # 8.5 - 0.6 - 0.6


def _section_header(doc: Document, text: str) -> None:
    """11pt bold UPPERCASE with a 1pt gray bottom rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "444444")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"


def _set_arial(run, size: int = 10, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _inline_bold(p, text: str, *, base_size: int = 10, italic: bool = False) -> None:
    """Append `text` to paragraph `p`, treating `**span**` as bold runs."""
    if not text:
        return
    parts = text.split("**")
    bold = False
    for i, segment in enumerate(parts):
        if i > 0:
            bold = not bold
        if not segment:
            continue
        run = p.add_run(segment)
        _set_arial(run, size=base_size, bold=bold, italic=italic)


def build_docx(resume: Resume, output_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    # ----- Header: Name + contact -----
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(resume.name)
    name_run.bold = True
    name_run.font.name = "Arial"
    name_run.font.size = Pt(18)
    name_run.font.spacing = Pt(2)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(2)
    contact_p.paragraph_format.space_after = Pt(2)
    _set_arial(contact_p.add_run(resume.contact_line), size=10)

    # ----- Professional Summary -----
    if resume.summary:
        _section_header(doc, "Professional Summary")
        sum_p = doc.add_paragraph()
        _inline_bold(sum_p, resume.summary, base_size=10)

    # ----- Professional Highlights -----
    if resume.highlights:
        _section_header(doc, "Professional Highlights")
        for h in resume.highlights:
            if not h:
                continue
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(2)
            _inline_bold(bp, h, base_size=10)

    # ----- Professional Experience -----
    if resume.experiences:
        _section_header(doc, "Professional Experience")
        for idx, exp in enumerate(resume.experiences):
            comp_p = doc.add_paragraph()
            if idx > 0:
                comp_p.paragraph_format.space_before = Pt(8)
            _set_arial(comp_p.add_run(exp.company), size=11, bold=True)

            # Meta line: italic title|location LEFT, dates RIGHT (tab stop)
            if exp.title or exp.location or exp.dates:
                meta_p = doc.add_paragraph()
                meta_p.paragraph_format.space_after = Pt(2)
                meta_p.paragraph_format.tab_stops.add_tab_stop(
                    Inches(CONTENT_WIDTH_INCHES), WD_TAB_ALIGNMENT.RIGHT
                )
                left_text = (
                    f"{exp.title} | {exp.location}"
                    if (exp.title and exp.location)
                    else (exp.title or exp.location)
                )
                _set_arial(meta_p.add_run(left_text), size=10, italic=True)
                if exp.dates:
                    meta_p.add_run("\t")
                    _set_arial(meta_p.add_run(exp.dates), size=10)

            # Optional descriptor (italic, full-width)
            if exp.descriptor:
                desc_p = doc.add_paragraph()
                desc_p.paragraph_format.space_after = Pt(3)
                _inline_bold(desc_p, exp.descriptor, base_size=10, italic=True)

            for bullet in exp.bullets:
                if not bullet:
                    continue
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)
                _inline_bold(bp, bullet, base_size=10)

    # ----- Selected Projects -----
    if resume.projects:
        _section_header(doc, "Selected Projects")
        for proj in resume.projects:
            name = (proj.name or "").strip()
            desc = (proj.description or "").strip()
            if not name and not desc:
                continue
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(3)
            if name:
                _set_arial(bp.add_run(name), size=10, bold=True)
            if name and desc:
                _set_arial(bp.add_run(": "), size=10)
            if desc:
                _inline_bold(bp, desc, base_size=10)

    # ----- Education and Certifications -----
    rows = list(resume.education) + list(resume.certifications)
    if rows:
        _section_header(doc, "Education and Certifications")
        for entry in rows:
            if not isinstance(entry, str) or not entry.strip():
                continue
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(2)
            _inline_bold(bp, entry, base_size=10)

    # ----- Core Technical Skills -----
    if resume.skills:
        _section_header(doc, "Core Technical Skills")
        for cat in resume.skills:
            if not cat.label or not cat.items:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(3)
            _set_arial(p.add_run(f"{cat.label}: "), size=10, bold=True)
            _set_arial(p.add_run(", ".join(cat.items)), size=10)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
