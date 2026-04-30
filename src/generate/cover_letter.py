"""JD-tailored cover letter generator.

Produces a 3-paragraph cover letter following the writing style rules from
the dheeraj-job-search skill: no em dashes, no en dashes in prose, direct
voice, numbers and named tools.
"""

from __future__ import annotations

import logging
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from ..llm import chat
from ..utils import OUTPUT_DIR, safe_filename_part, scrub_dashes
from . import mirror_to_public

log = logging.getLogger(__name__)

SYSTEM = """You are writing a cover letter for Dheeraj Sampath, an Engineering Manager based in Austin, TX with 14+ years of frontend engineering and engineering leadership experience.

LOCKED EXPERIENCE
- Equifax, Engineering Lead, Austin TX, Aug 2022 to Present (front-end architecture, micro-frontends, Module Federation, 10+ launches, 50% scalability gain, 35% delivery efficiency, 25% build cycle reduction)
- Midigator, Front End Architect, Austin TX (2022, embeddable React UI, 35% time-to-market reduction)
- TA Digital, UI Architect (2018-2022, 400% page load improvement, AI scalable solutions)
- 14+ years total. MBA Digital Entrepreneurship, B.Tech Computer Science.

WRITING STYLE RULES (HARD)
- No em dashes (—) anywhere
- No en dashes (–) in prose
- Replace dashes with period, comma, colon, parentheses, or semicolon
- Direct voice, short sentences
- Quantify when possible (specific numbers from the locked experience)
- No flattery about the company
- No "I am writing to apply for"
- Open with a specific hook tied to the role
- Close with a clear next step

STRUCTURE
- Paragraph 1 (3-4 sentences): hook tied to the specific role + your strongest relevant signal
- Paragraph 2 (4-6 sentences): two concrete achievements with numbers that map to the JD's top 2-3 priorities
- Paragraph 3 (2-3 sentences): why this company specifically + clear next step

OUTPUT
Return ONLY the cover letter text. No greeting line, no signature, no JSON, no markdown. Just three paragraphs separated by blank lines."""


def expected_cover_letter_path(jd_title: str, jd_company: str) -> Path:
    safe_t = safe_filename_part(jd_title)
    safe_c = safe_filename_part(jd_company)
    return OUTPUT_DIR / f"CoverLetter_Dheeraj_Sampath_{safe_t}_{safe_c}.docx"


def generate_cover_letter(
    jd_title: str, jd_company: str, jd_text: str, model: str | None = None
) -> tuple[Path, str]:
    model = model or os.environ.get("GENERATION_MODEL", "anthropic/claude-sonnet-4.5")

    user = f"""TARGET TITLE: {jd_title}
TARGET COMPANY: {jd_company}

JOB DESCRIPTION:
{jd_text[:8000] if jd_text else "(JD not available - write conservatively)"}

Write the cover letter now."""

    text = chat(system=SYSTEM, user=user, model=model, max_tokens=1500).strip()
    text = scrub_dashes(text)

    # Build .docx — match the resume's Arial / minimal-sleek treatment so the
    # two artifacts feel like a coordinated set.
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.25

    # Header — matches resume name treatment (centered, ALL CAPS, letter-spaced)
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run("DHEERAJ SAMPATH")
    name_run.bold = True
    name_run.font.name = "Arial"
    name_run.font.size = Pt(16)
    name_run.font.spacing = Pt(2)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = contact_p.add_run("Austin, TX | 248-873-8929 | dheerajsampath@proton.me")
    cr.font.name = "Arial"
    cr.font.size = Pt(10)

    # Subject line
    subj_p = doc.add_paragraph()
    subj_p.paragraph_format.space_before = Pt(14)
    sr = subj_p.add_run(f"Re: {jd_title} at {jd_company}")
    sr.bold = True
    sr.font.name = "Arial"
    sr.font.size = Pt(11)

    doc.add_paragraph("")

    for para in text.split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("")
    doc.add_paragraph("Best,")
    doc.add_paragraph("Dheeraj Sampath")

    output_path = expected_cover_letter_path(jd_title, jd_company)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    public_path = mirror_to_public(output_path)
    if public_path:
        log.info("cover letter mirrored to %s", public_path)

    return output_path, text
