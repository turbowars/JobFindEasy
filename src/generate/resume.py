"""JD-tailored resume generator.

Calls Claude with the dheeraj-resume-generator skill content embedded as
system prompt, asks for structured JSON (headline + bullets + skills), then
assembles a .docx via python-docx following the formatting rules in the skill.

The skill file path is read at runtime so updates to the skill propagate.
"""
from __future__ import annotations

import hashlib
import json
import logging
import os
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..llm import chat
from ..enrichment.ats_match import extract_keywords, match_keywords
from ..enrichment.hr_score import hr_simulate
from . import mirror_to_public

log = logging.getLogger(__name__)

# Prefer the in-project copy so the skill travels with the repo.
PROJECT_ROOT = Path(__file__).parent.parent.parent
SKILL_PATH_CANDIDATES = [
    PROJECT_ROOT / "config" / "skills" / "dheeraj-resume-generator" / "SKILL.md",
    Path.home() / ".claude" / "skills" / "dheeraj-resume-generator" / "SKILL.md",
    Path("/mnt/skills/user/dheeraj-resume-generator/SKILL.md"),
]

OUTPUT_DIR = PROJECT_ROOT / "data" / "exports"


def _load_skill() -> str:
    """Load SKILL.md and inline its references/ files so the model has the full
    context (the SKILL document references `references/profile-{ic,em}.md` by
    relative path, which the model cannot resolve on its own)."""
    for skill_path in SKILL_PATH_CANDIDATES:
        if skill_path.exists():
            content = skill_path.read_text()
            refs_dir = skill_path.parent / "references"
            if refs_dir.is_dir():
                for ref in sorted(refs_dir.glob("*.md")):
                    content += (
                        f"\n\n============================================================\n"
                        f"REFERENCE FILE: references/{ref.name}\n"
                        f"============================================================\n"
                        f"{ref.read_text()}"
                    )
            return content
    log.warning("resume skill file not found at any candidate path; using built-in fallback")
    return _FALLBACK_SKILL


_FALLBACK_SKILL = """You are generating a resume for Dheeraj Sampath, an Engineering Manager based in Austin, TX.

LOCKED EXPERIENCE (do not invent or alter):
- Equifax, Engineering Lead, Austin TX, Aug 2022 to Present
- Midigator, Front End Architect, Austin TX, Feb 2022 to Aug 2022
- TA Digital, UI Architect, Minneapolis MN, Feb 2018 to Feb 2022
- NowFloats Technologies, Principal Engineer, India, Mar 2017 to Feb 2018
- Deloitte Digital Studio, Senior Engineer, Mumbai, Aug 2014 to Mar 2017
- Neudesic, Consultant, Hyderabad, May 2012 to Aug 2014

EDUCATION
- MBA Digital Entrepreneurship, Strayer University
- B.Tech Computer Science, Mahatma Gandhi Institute of Technology

CONTACT (EM profile)
- Austin, TX
- 248-873-8929
- dheerajsampath@proton.me
- linkedin.com/in/evolvingdx
- dheerajsampath.framer.website
- github.com/turbowars

WRITING RULES
- ZERO em dashes (—) anywhere in the resume — including education and project entries.
- ZERO en dashes (–) anywhere; use commas or " - " (hyphen with spaces) instead.
- ZERO IC / "Individual Contributor" track markers in any visible field.
- Quantify every bullet (number, percent, scope, named tool).

Return JSON only with the structure described in the user message."""


SYSTEM_TEMPLATE = """{skill_content}

============================================================
OUTPUT REQUIREMENT
You must return ONLY a single JSON object. No markdown fences, no commentary.
The JSON must have this exact shape:

{{
  "track": "EM" or "IC",
  "headline": "Slot1  |  Slot2  |  Slot3",
  "summary": "3-4 sentence professional summary, no em dashes",
  "skills": [
    {{"label": "Languages", "items": ["TypeScript", "JavaScript (ES2023)", "Go", ...]}},
    {{"label": "Frontend", "items": ["React 18", "Next.js 14", ...]}},
    ...
  ],
  "experience": [
    {{
      "company": "Equifax",
      "title": "the title to print, with mirror parenthetical if EM",
      "location": "Austin, TX",
      "dates": "Aug 2022 - Present",
      "bullets": [
        "bullet 1 with **bold metric phrase** for emphasis",
        "bullet 2 ..."
      ]
    }},
    ... more roles in reverse chronological order
  ],
  "projects": [
    {{"name": "AI Code Review Agent", "description": "Next.js + Claude API + AST parsing. ..."}},
    ... 3-5 selected projects, ONLY for IC track. EM track: omit or empty array.
  ],
  "education": [
    "M.B.A., Digital Entrepreneurship - Strayer University, Herndon, VA",
    "B.Tech, Computer Science - Mahatma Gandhi Institute of Technology, India",
    "Certified Scrum Master, Scrum Alliance | Certified Usability Analyst, HFI"
  ],
  "tailoring_report": {{
    "profile_used": "EM" or "IC",
    "headline": "the 3-slot headline",
    "title_mirror": "Equifax line title",
    "keyword_match": "X/10 top JD keywords hit",
    "priorities_addressed": ["JD priority 1 -> bullet that covers it", ...],
    "missing_signals": ["anything the JD wants that the candidate's background doesn't cover"]
  }}
}}

FORMATTING DIRECTIVES (follow exactly):
- Skills are LABELED CATEGORIES (e.g. "Languages", "Frontend", "Backend and APIs",
  "AI / LLM Engineering", "Cloud and Infrastructure", "Architecture"). Each
  category groups 6-15 short items. NOT bullets, NOT prose paragraphs.
- Experience bullets MAY use **markdown bold** to emphasize metric phrases
  ("**cutting build cycles 25%**", "**MCP-Powered Dev Assistant**"). Use bold
  selectively — 1-2 emphasis spans per bullet at most. Bold ONLY the
  measurable result or a named tool/system, never filler words.
- Education entries use ` - ` (single hyphen with spaces) to separate degree
  from institution; keep one entry per line. Never em dashes.

HEADLINE RULES (track-conditional):
- IC track: headline is THREE slots separated by `  |  ` (two spaces, pipe,
  two spaces). Slot 1 = JD target title (exact string). Slots 2 and 3 = two
  closely-related title alternates a recruiter would consider equivalent.
- **EM track: headline is exactly ONE slot — the JD target title, verbatim.
  No pipes, no alternates, no parentheticals.**

NEVER OUTPUT THESE TOKENS ANYWHERE IN THE RESUME:
- The strings "IC", "EM", "(IC)", "(EM)", "IC track", "EM track",
  "Individual Contributor", or any other track marker — in ANY field
  (headline, summary, role title, bullet, skill, project, education).
  The track is INTERNAL metadata; it must never reach a visible field.
- Em dashes (—) and en dashes (–) anywhere. Use commas or " - " (hyphen
  with spaces) instead. This is a hard rule — no exceptions, including
  education entries.
- The target company's name embedded inside a mirror role title. The JD
  target role MUST NOT contain the company you are applying to. Example
  WRONG: if applying to Coinbase for "Engineering Manager, Legend", do NOT
  emit a role titled "Engineering Manager, Legend" at Equifax — the mirror
  role at Equifax should describe the role itself, not the target team or
  product. RIGHT: drop any team/product qualifier that ONLY exists at the
  target company. Keep generic role qualifiers ("Frontend", "Platform",
  "Identity") only if they describe a transferable scope.

CANONICAL COMPANY NAMES IN EXPERIENCE BLOCK (always use these exact spellings):
- NowFloats Technologies / NowFloats Technologies Ltd. → "nowfloats" (lowercase)
- Equifax → "Equifax"
- Midigator → "Midigator"
- TA Digital → "TA Digital"
- Deloitte Digital Studio → "Deloitte Digital Studio"
- Neudesic → "Neudesic"

Final check before returning:
- ZERO em dashes (—) and ZERO en dashes (–) anywhere — every separator is a
  comma or " - " (hyphen with spaces). Education and project entries are
  not exceptions.
- ZERO occurrences of "IC", "EM", "IC track", "EM track", or
  "Individual Contributor" in any visible field.
- Headline slot 1 = JD target title, exact string.
- For EM track: headline contains exactly ONE slot (no pipes).
- Summary line 1 starts with the same noun phrase as headline slot 1.
- No mirror role title containing the target company's name.
- Output is valid JSON parseable by Python's json.loads.
"""


def _generate_structured(
    model: str,
    jd_title: str,
    jd_company: str,
    jd_text: str,
    must_cover: Optional[dict] = None,
    retry_feedback: Optional[str] = None,
) -> dict:
    user_parts = [
        f"TARGET TITLE: {jd_title}",
        f"TARGET COMPANY: {jd_company}",
        "",
        "JOB DESCRIPTION:",
        jd_text[:10000] if jd_text else "(JD text not available - score conservatively)",
        "",
    ]
    if must_cover and any(must_cover.get(t) for t in ("required", "preferred", "soft")):
        user_parts.append("MUST-COVER ATS KEYWORDS (use these exact spellings where natural):")
        for tier in ("required", "preferred", "soft"):
            items = must_cover.get(tier) or []
            if items:
                user_parts.append(f"- {tier.capitalize()}: {', '.join(items)}")
        user_parts.append("")
    if retry_feedback:
        user_parts.append("RETRY FEEDBACK (the previous attempt scored low — fix these):")
        user_parts.append(retry_feedback)
        user_parts.append("")
    user_parts.append("Generate the tailored resume JSON now.")
    user = "\n".join(user_parts)

    skill = _load_skill()
    sys_prompt = SYSTEM_TEMPLATE.format(skill_content=skill)

    # cache_system=True: the 54KB skill content + template is the dominant
    # input cost (~14k tokens). Cached after the first call within 5 min,
    # cutting Sonnet input cost ~40% on subsequent generations.
    #
    # max_tokens=8192: a complete IC-track resume JSON with 7 roles, 5-6
    # bullets each, 4 projects, structured skills can run ~4-5k tokens.
    # 4096 was right at the edge and would occasionally truncate the final
    # role, causing the HR scorer to flag "resume cuts off mid-sentence".
    text = chat(
        system=sys_prompt, user=user, model=model,
        max_tokens=8192, cache_system=True,
    ).strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.startswith("json"):
            text = text[4:]
        text = text.strip()
    return json.loads(text)


def _scrub_dashes(s: str) -> str:
    """Replace every em dash (—) and en dash (–) with a hyphen-with-spaces
    or a plain comma where adjacent whitespace makes the hyphen feel wrong.
    Applied universally — there are no fields where em dashes are wanted.
    """
    if not s:
        return s
    # ` — ` and ` – ` (separator forms with spaces) become ` - `
    s = re.sub(r"\s*[—–]\s*", " - ", s)
    # Collapse double spaces introduced by the substitution
    s = re.sub(r"  +", " ", s)
    return s.strip()


# Belt-and-suspenders normalization for output the LLM might still drift on.
# Catches every track-marker form we've seen the model emit:
#   "(IC)" / "(EM)" / "(ic)" / "(em)"
#   "(Individual Contributor)"
#   "IC track" / "EM track" / "IC Track" / "EM Track"
#   " — IC", " - EM", ", IC", ": EM"  (suffix tags after a separator)
#   "Individual Contributor" (bare)
_TRACK_TAG_PATTERNS = [
    re.compile(r"\s*\((?:IC|EM|Individual\s+Contributor)\)\s*", re.IGNORECASE),
    re.compile(r"\s*[—\-,:]\s*(?:IC|EM)(?:\s+[Tt]rack)?(?=\s|$|[.,;:|])"),
    re.compile(r"\b(?:IC|EM)\s+[Tt]rack\b"),
    re.compile(r"\bIndividual\s+Contributor(?:\s+[Tt]rack)?\b", re.IGNORECASE),
]

# Map of "wrong" company strings the LLM sometimes emits → canonical form.
# Applied to BOTH the company name field and any embedded mention in a title.
_COMPANY_CANONICAL = [
    ("NowFloats Technologies Ltd.", "nowfloats"),
    ("NowFloats Technologies Ltd", "nowfloats"),
    ("NowFloats Technologies, Ltd.", "nowfloats"),
    ("NowFloats Technologies", "nowfloats"),
    ("NowFloats", "nowfloats"),
]


def _strip_track_tags(s: str) -> str:
    if not s:
        return s
    for pat in _TRACK_TAG_PATTERNS:
        s = pat.sub("", s)
    # Collapse whitespace introduced by the substitutions.
    return re.sub(r"\s{2,}", " ", s).strip()


def _canonicalize_company(s: str) -> str:
    if not s:
        return s
    for wrong, right in _COMPANY_CANONICAL:
        s = s.replace(wrong, right)
    return s


def _strip_company_from_title(title: str, target_company: str) -> str:
    """Remove the target company's name from a mirror role title.

    Patterns we strip:
      - "Engineering Manager, <Company>"  → "Engineering Manager"
      - "Engineering Manager — <Company>" → "Engineering Manager"
      - "<Company> Engineering Manager"   → "Engineering Manager"
    Only the trailing/leading occurrence is removed; embedded scope words
    that happen to share letters with the company are left alone.
    """
    if not title or not target_company:
        return title
    tc = target_company.strip()
    if not tc:
        return title
    # Trailing patterns
    for sep in (", ", " — ", " - ", " | "):
        suffix = f"{sep}{tc}"
        if title.endswith(suffix):
            return title[: -len(suffix)].strip()
    # Leading patterns
    prefix = f"{tc} "
    if title.startswith(prefix):
        return title[len(prefix):].strip()
    return title


def _canonicalize_payload(payload: dict, jd_company: str, track: str) -> dict:
    """Defensive post-processing: enforce the formatting rules from the prompt
    even when the LLM drifts. Idempotent — safe to run twice.
    """
    # Strip track tags everywhere they might appear.
    if "headline" in payload:
        payload["headline"] = _strip_track_tags(payload["headline"] or "")
    if "summary" in payload:
        payload["summary"] = _strip_track_tags(payload["summary"] or "")
    for role in payload.get("experience", []) or []:
        role["company"] = _canonicalize_company(_strip_track_tags(role.get("company") or ""))
        title = _strip_track_tags(role.get("title") or "")
        # Mirror title shouldn't contain the target company's name.
        title = _strip_company_from_title(title, jd_company)
        role["title"] = title
        role["bullets"] = [
            _strip_track_tags(b or "") for b in (role.get("bullets") or [])
        ]
    payload["education"] = [
        _strip_track_tags(_canonicalize_company(e or ""))
        for e in (payload.get("education") or [])
    ]
    for proj in payload.get("projects", []) or []:
        if isinstance(proj, dict):
            proj["name"] = _strip_track_tags(proj.get("name") or "")
            proj["description"] = _strip_track_tags(proj.get("description") or "")

    # EM headline collapses to a single slot (the JD target title, slot 1).
    if (track or "").upper() == "EM" and payload.get("headline"):
        slots = [s.strip() for s in payload["headline"].split("|") if s.strip()]
        if len(slots) > 1:
            payload["headline"] = slots[0]

    return payload


def _scrub_payload(payload: dict) -> dict:
    """Strip em / en dashes from every renderable field. No field is exempt
    — em dashes don't appear in the resume regardless of context.
    """
    if "headline" in payload:
        payload["headline"] = _scrub_dashes(payload.get("headline") or "")
    if "summary" in payload:
        payload["summary"] = _scrub_dashes(payload.get("summary") or "")
    for role in payload.get("experience", []) or []:
        role["title"] = _scrub_dashes(role.get("title") or "")
        role["company"] = _scrub_dashes(role.get("company") or "")
        role["location"] = _scrub_dashes(role.get("location") or "")
        # Date ranges intentionally use ` - ` (hyphen) per the prompt rules,
        # but if the LLM ever emits em/en dashes in dates, scrub them too.
        role["dates"] = _scrub_dashes(role.get("dates") or "")
        bullets = role.get("bullets") or []
        role["bullets"] = [_scrub_dashes(b) for b in bullets]
    payload["education"] = [
        _scrub_dashes(e or "") for e in (payload.get("education") or [])
    ]
    for proj in payload.get("projects", []) or []:
        if isinstance(proj, dict):
            proj["name"] = _scrub_dashes(proj.get("name") or "")
            proj["description"] = _scrub_dashes(proj.get("description") or "")
    for skill_block in payload.get("skills", []) or []:
        if isinstance(skill_block, dict):
            skill_block["label"] = _scrub_dashes(skill_block.get("label") or "")
            items = skill_block.get("items") or []
            skill_block["items"] = [_scrub_dashes(i or "") for i in items]
    return payload


def _flatten_payload(payload: dict) -> str:
    """Concatenate all renderable text into a single searchable blob for ATS
    keyword matching and HR scoring. Avoids round-tripping through the .docx
    + mammoth.
    """
    parts: list[str] = []
    if payload.get("headline"):
        parts.append(payload["headline"])
    if payload.get("summary"):
        parts.append(payload["summary"])
    for cat in payload.get("skills", []) or []:
        if isinstance(cat, dict):
            label = cat.get("label", "")
            items = cat.get("items", []) or []
            parts.append(f"{label}: {', '.join(items)}")
        elif isinstance(cat, str):
            parts.append(cat)
    for role in payload.get("experience", []) or []:
        parts.append(f"{role.get('company','')} {role.get('title','')} {role.get('location','')}")
        for b in role.get("bullets", []) or []:
            # Strip ** markers so partial_ratio sees clean text
            parts.append(b.replace("**", ""))
    for proj in payload.get("projects", []) or []:
        if isinstance(proj, dict):
            parts.append(f"{proj.get('name','')} {proj.get('description','')}")
    for edu in payload.get("education", []) or []:
        if isinstance(edu, str):
            parts.append(edu)
    return "\n".join(parts)


def _build_retry_feedback(missing: dict, weakest_areas: list) -> str:
    lines = []
    flat_missing = []
    for tier in ("required", "preferred", "soft"):
        items = missing.get(tier) or []
        if items:
            flat_missing.extend(items)
    if flat_missing:
        lines.append(f"Missing keywords (weave these in naturally): {', '.join(flat_missing[:20])}.")
    if weakest_areas:
        lines.append("HR concerns: " + "; ".join(weakest_areas[:5]) + ".")
    if not lines:
        lines.append("Improve specificity and metrics in bullets; reduce buzzword density.")
    return " ".join(lines)


CONTACT_LINE = (
    "Austin, TX | 248-873-8929 | dheerajsampath@proton.me | "
    "linkedin.com/in/evolvingdx | dheerajsampath.com"
)
NAME_TEXT = "DHEERAJ SAMPATH"
CONTENT_WIDTH_INCHES = 7.3  # 8.5 - 0.6 - 0.6 = 7.3" between margins


def _add_section_header(doc: Document, text: str) -> None:
    """Section header: 11pt bold UPPERCASE with a 1pt gray bottom rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")  # 1pt rule (sz is in 1/8 pt units)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "444444")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"


def _add_inline_bold(p, text: str, *, base_size: int = 10, italic: bool = False) -> None:
    """Append `text` to paragraph `p`, treating `**span**` markers as bold runs.

    Splits on `**` boundaries; alternating segments toggle bold. Any leftover
    odd `**` is treated as plain text. Empty segments are skipped.
    """
    if not text:
        return
    parts = text.split("**")
    bold = False
    for i, segment in enumerate(parts):
        # Toggle bold at each boundary except the first
        if i > 0:
            bold = not bold
        if not segment:
            continue
        run = p.add_run(segment)
        run.font.name = "Arial"
        run.font.size = Pt(base_size)
        run.bold = bold
        if italic:
            run.italic = True


def _set_run_arial(run, size: int = 10, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _build_docx(payload: dict, output_path: Path) -> None:
    """Assemble a .docx matching the reference IC Staff resume layout.

    Single column, all-black, sans-serif Arial. Section headers are bold
    UPPERCASE with a 1pt gray bottom rule. Skills render as labeled
    paragraphs (`<b>Label:</b> items`). Experience uses a tab-stopped meta
    line so dates right-align. Bullets support inline `**bold**` markers
    for selective metric emphasis.
    """
    doc = Document()

    # Tight margins so we get more room — reference is dense.
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    # ----- Header block -----
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(NAME_TEXT)
    name_run.bold = True
    name_run.font.name = "Arial"
    name_run.font.size = Pt(18)
    name_run.font.spacing = Pt(2)  # letter-spacing

    # Sub-headline (3-slot)
    headline_text = (payload.get("headline") or "").strip()
    if headline_text:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_before = Pt(2)
        _set_run_arial(sub_p.add_run(headline_text), size=11)

    # Contact
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(2)
    _set_run_arial(contact_p.add_run(CONTACT_LINE), size=10)

    # ----- Summary -----
    if payload.get("summary"):
        _add_section_header(doc, "Summary")
        sum_p = doc.add_paragraph()
        _add_inline_bold(sum_p, payload["summary"], base_size=10)

    # ----- Skills (labeled categories) -----
    skills = payload.get("skills") or []
    if skills:
        _add_section_header(doc, "Core Technical Skills")
        for cat in skills:
            if isinstance(cat, dict):
                label = (cat.get("label") or "").strip()
                items = cat.get("items") or []
                if not label or not items:
                    continue
                items_text = ", ".join(str(x) for x in items)
                p = doc.add_paragraph()
                # Hanging indent so wrap-lines align under items
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.first_line_indent = Inches(-0.18)
                p.paragraph_format.space_after = Pt(3)
                _set_run_arial(p.add_run(f"{label}: "), size=10, bold=True)
                _set_run_arial(p.add_run(items_text), size=10)
            elif isinstance(cat, str):
                # Backwards-compat with old "Label: items" string form
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.first_line_indent = Inches(-0.18)
                p.paragraph_format.space_after = Pt(3)
                if ": " in cat:
                    label, rest = cat.split(": ", 1)
                    _set_run_arial(p.add_run(f"{label}: "), size=10, bold=True)
                    _set_run_arial(p.add_run(rest), size=10)
                else:
                    _set_run_arial(p.add_run(cat), size=10)

    # ----- Experience -----
    experience = payload.get("experience") or []
    if experience:
        _add_section_header(doc, "Professional Experience")
        for idx, role in enumerate(experience):
            company = (role.get("company") or "").strip()
            title = (role.get("title") or "").strip()
            location = (role.get("location") or "").strip()
            dates = (role.get("dates") or "").strip()

            # Company line — bold, slight breath above
            comp_p = doc.add_paragraph()
            if idx > 0:
                comp_p.paragraph_format.space_before = Pt(8)
            _set_run_arial(comp_p.add_run(company), size=11, bold=True)

            # Meta line — italic title|location LEFT, dates RIGHT (tab-stopped)
            if title or location or dates:
                meta_p = doc.add_paragraph()
                meta_p.paragraph_format.space_after = Pt(2)
                # Right tab stop at content width
                meta_p.paragraph_format.tab_stops.add_tab_stop(
                    Inches(CONTENT_WIDTH_INCHES), WD_TAB_ALIGNMENT.RIGHT
                )
                left_text = f"{title} | {location}" if (title and location) else (title or location)
                _set_run_arial(meta_p.add_run(left_text), size=10, italic=True)
                if dates:
                    meta_p.add_run("\t")
                    _set_run_arial(meta_p.add_run(dates), size=10)

            # Bullets
            for bullet in role.get("bullets") or []:
                if not bullet:
                    continue
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)
                _add_inline_bold(bp, bullet, base_size=10)

    # ----- Selected Projects (IC track) -----
    projects = payload.get("projects") or []
    if projects:
        _add_section_header(doc, "Selected Projects")
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            name = (proj.get("name") or "").strip()
            desc = (proj.get("description") or "").strip()
            if not name and not desc:
                continue
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(3)
            if name:
                _set_run_arial(bp.add_run(name), size=10, bold=True)
            if name and desc:
                _set_run_arial(bp.add_run(" — "), size=10)
            if desc:
                _add_inline_bold(bp, desc, base_size=10)

    # ----- Education -----
    education = payload.get("education") or []
    if education:
        _add_section_header(doc, "Education and Certifications")
        for edu in education:
            if not isinstance(edu, str) or not edu.strip():
                continue
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(2)
            _add_inline_bold(bp, edu, base_size=10)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _safe_loc_suffix(location: str) -> str:
    """Build the location suffix for the resume filename. Truncates at 30
    chars for readability, but appends a 6-char content hash when truncation
    happens so two distinct long locations cannot collide on the same prefix
    (e.g. "New York City, New York State..." vs "New York City, New York Store...").
    """
    if not location:
        return ""
    raw = re.sub(r"[^A-Za-z0-9]+", "_", location).strip("_")
    if not raw:
        return ""
    if len(raw) <= 30:
        return f"_{raw}"
    digest = hashlib.md5(location.encode()).hexdigest()[:6]
    return f"_{raw[:23]}_{digest}"


def expected_resume_path(jd_title: str, jd_company: str, location: str = "") -> Path:
    """Canonical filename `generate_resume` will write to (always includes
    the location suffix when one is provided). When the location string is
    long, a short content hash is appended to prevent collisions between two
    different long locations whose first 30 chars happen to match."""
    safe_t = re.sub(r"[^A-Za-z0-9]+", "_", jd_title).strip("_")
    safe_c = re.sub(r"[^A-Za-z0-9]+", "_", jd_company).strip("_")
    suffix = _safe_loc_suffix(location)
    return OUTPUT_DIR / f"Dheeraj_Sampath_{safe_t}_{safe_c}{suffix}.docx"


def existing_resume_path(jd_title: str, jd_company: str, location: str = "") -> Path:
    """Returns the existing resume on disk for this job, accepting either the
    new (location-suffixed) or the legacy (no-suffix) filename format. Used
    by lookup paths (UI preview, autogen skip-if-exists) so we don't miss
    files that were generated before the location-aware naming was added.

    If neither variant exists, falls back to the canonical (with-location)
    path so callers can still write to `.exists()` and get False.
    """
    canonical = expected_resume_path(jd_title, jd_company, location)
    if canonical.exists():
        return canonical
    legacy = expected_resume_path(jd_title, jd_company, "")
    if legacy.exists():
        return legacy
    return canonical


ATS_RETRY_THRESHOLD = 75
HR_RETRY_THRESHOLD = 70


def today_resume_count() -> int:
    """Count resumes already generated TODAY (by file mtime).

    Used as a cheap rate-limiter for the per-day generation cap. Reading the
    file system is fast and self-correcting — if the user manually deletes
    files, the count drops accordingly. No DB schema needed.
    """
    from datetime import date, datetime
    today = date.today()
    n = 0
    if not OUTPUT_DIR.exists():
        return 0
    for p in OUTPUT_DIR.glob("Dheeraj_Sampath_*.docx"):
        try:
            if datetime.fromtimestamp(p.stat().st_mtime).date() == today:
                n += 1
        except Exception:
            continue
    return n


def daily_cap_reached() -> bool:
    """True iff today's generation count is at or above the per-day cap.

    Cap is read from $AUTO_RESUME_CAP_PER_DAY (default 20). Check this
    BEFORE submitting a generation in any auto path (autoscrape, CLI,
    autogen-on-score). Manual UI generations are allowed past the cap —
    those are user-initiated, the cap is just a runaway-cost guard.
    """
    cap = int(os.environ.get("AUTO_RESUME_CAP_PER_DAY", "20"))
    return today_resume_count() >= cap


def generate_resume(
    jd_title: str, jd_company: str, jd_text: str,
    model: Optional[str] = None, location: str = "",
) -> tuple[Path, dict]:
    """Generate a tailored resume with the dual-scored auto-iterating pipeline.

    Pipeline:
      A. Haiku extract JD keywords (cached per JD)
      B. Sonnet generate resume payload, with keyword list as must-cover hint
      C. rapidfuzz keyword match → coverage %
      D. Haiku HR simulation → recruiter-perspective score + weakest areas
      E. If either score < threshold: ONE retry of B with explicit feedback
      F. Build .docx, mirror to ~/Public, write a .scores.json sidecar

    Returns (path_to_docx, tailoring_report_dict). The tailoring report
    includes a `scores` key with both ATS and HR results.
    """
    model = model or os.environ.get("GENERATION_MODEL", "anthropic/claude-sonnet-4.5")
    cache_key = f"{jd_company}|{jd_title}"

    # [A] pre-flight keyword extraction
    keywords = extract_keywords(jd_text or "", cache_key)
    log.info(
        "[ats_match] keywords for %s: %d required / %d preferred / %d soft",
        cache_key,
        len(keywords.get("required", [])),
        len(keywords.get("preferred", [])),
        len(keywords.get("soft", [])),
    )

    # [B] generate
    payload = _generate_structured(
        model, jd_title, jd_company, jd_text, must_cover=keywords
    )
    payload = _scrub_payload(payload)
    payload = _canonicalize_payload(payload, jd_company, payload.get("track") or "")

    # [C] keyword match
    resume_text = _flatten_payload(payload)
    match = match_keywords(resume_text, keywords)
    # [D] HR simulation
    hr = hr_simulate(jd_title, jd_company, jd_text, resume_text)
    log.info(
        "[scores] ats=%s hr=%s for %s",
        match.get("match_pct"), hr.get("hr_score"), cache_key,
    )

    # [E] auto-iterate ONCE if either score is below threshold
    retried = False
    if (match.get("match_pct", 0) < ATS_RETRY_THRESHOLD) or (
        hr.get("hr_score", 100) < HR_RETRY_THRESHOLD
    ):
        feedback = _build_retry_feedback(
            match.get("missing", {}),
            hr.get("weakest_areas", []),
        )
        log.info("[retry] regenerating once with feedback: %s", feedback[:200])
        try:
            payload = _generate_structured(
                model, jd_title, jd_company, jd_text,
                must_cover=keywords, retry_feedback=feedback,
            )
            payload = _scrub_payload(payload)
            payload = _canonicalize_payload(payload, jd_company, payload.get("track") or "")
            resume_text = _flatten_payload(payload)
            match = match_keywords(resume_text, keywords)
            hr = hr_simulate(jd_title, jd_company, jd_text, resume_text)
            retried = True
            log.info(
                "[retry] post-retry scores: ats=%s hr=%s",
                match.get("match_pct"), hr.get("hr_score"),
            )
        except Exception as e:
            log.warning("[retry] regeneration failed; keeping first attempt: %s", e)

    # [F] build .docx + sidecar + mirror
    output_path = expected_resume_path(jd_title, jd_company, location)
    _build_docx(payload, output_path)

    scores = {
        "ats_match": match,
        "hr": hr,
        "retried": retried,
        "keywords": keywords,
    }
    sidecar = output_path.with_suffix(".scores.json")
    try:
        sidecar.write_text(json.dumps(scores, indent=2))
    except Exception as e:
        log.warning("could not write scores sidecar: %s", e)

    public_path = mirror_to_public(output_path)
    if public_path:
        log.info("resume mirrored to %s", public_path)
    public_sidecar = mirror_to_public(sidecar) if sidecar.exists() else None
    if public_sidecar:
        log.info("scores sidecar mirrored to %s", public_sidecar)

    tailoring = payload.get("tailoring_report") or {}
    tailoring["scores"] = scores
    return output_path, tailoring


def autogen_resume_if_missing(
    jd_title: str, jd_company: str, jd_text: str, location: str = ""
) -> Optional[Path]:
    """Generate a resume only if one for this title+company+location doesn't exist.

    Returns the resume path (existing or freshly generated), or None on failure
    or when the daily cap has been reached. Safe to call from any thread.
    Accepts both new (location-suffixed) and legacy (no-suffix) filename
    formats when checking for existence.
    """
    existing = existing_resume_path(jd_title, jd_company, location)
    if existing.exists():
        log.info("resume exists, skipping autogen: %s", existing.name)
        return existing
    if daily_cap_reached():
        cap = int(os.environ.get("AUTO_RESUME_CAP_PER_DAY", "20"))
        log.info(
            "daily resume cap reached (%d generations today >= %d) — skipping autogen for %s @ %s",
            today_resume_count(), cap, jd_title, jd_company,
        )
        return None
    try:
        path, _ = generate_resume(jd_title, jd_company, jd_text, location=location)
        log.info("auto-generated resume: %s", path.name)
        return path
    except Exception as e:
        log.warning(
            "autogen resume failed for %s @ %s (%s): %s",
            jd_title, jd_company, location, e,
        )
        return None
